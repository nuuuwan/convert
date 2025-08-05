from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import Log

from convert.core.AbstractDoc import AbstractDoc
from convert.core.Paragraph import Paragraph

log = Log("DocXDoc")


class DocXDoc(AbstractDoc):

    @staticmethod
    def parse_docx_paragraph(docx_paragraph) -> Paragraph:
        if docx_paragraph.style.name.startswith("Heading"):
            level = int(docx_paragraph.style.name[-1])
            return Paragraph(f"h{level}", docx_paragraph.text)
        return Paragraph("p", docx_paragraph.text)

    @classmethod
    def from_file(cls, file_path: str) -> None:
        doc = Document(file_path)
        paragraphs = []
        for docx_paragraph in doc.paragraphs:
            if docx_paragraph.text.strip():
                paragraph = DocXDoc.parse_docx_paragraph(docx_paragraph)
                paragraphs.append(paragraph)
        doc = cls(paragraphs)
        doc.clean()
        return doc

    def to_file(self, file_path: str) -> None:
        doc = Document()

        for paragraph in self.paragraphs:
            if paragraph.tag.startswith("h"):
                level = int(paragraph.tag[1])
                doc.add_heading(paragraph.text, level=level)
            else:
                p = doc.add_paragraph(paragraph.text)
                if paragraph.text == "...":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(file_path)
